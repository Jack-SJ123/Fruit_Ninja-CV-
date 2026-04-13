"""
Generate two Phase 1 Word documents for CV Final Project.
Version A: Hand-Gesture Controlled Fruit-Slicing Game
Version B: Brain Tumor Classification (Scenario 1)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="4472C4", sz="4"):
    """Set borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_formatted_table(doc, headers, rows, header_color="2E4057"):
    """Add a nicely formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
        set_cell_borders(cell, color=header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(10)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F2F6FA")
            set_cell_borders(cell, color="B0BEC5", sz="2")

    return table


def setup_document(title, subtitle, group_info):
    """Create a new document with cover-page-style header."""
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Configure heading styles
    for level, (size, color_hex) in enumerate(
        [(18, "1B3A5C"), (14, "2E4057"), (12, "3D5A80")], start=1
    ):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        )
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(6)

    # ── Title block ──
    doc.add_paragraph()  # spacer

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x3D, 0x5A, 0x80)

    doc.add_paragraph()  # spacer

    p_group = doc.add_paragraph()
    p_group.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_group.add_run(group_info)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Add a line break before content
    doc.add_paragraph()
    doc.add_page_break()

    return doc


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        p.add_run(text)


# ═══════════════════════════════════════════════════════════════════════════
# VERSION A: Hand-Gesture Controlled Fruit-Slicing Game
# ═══════════════════════════════════════════════════════════════════════════

def create_version_a(output_dir):
    group = (
        "Group 9\n"
        "Angel Daniel Bustamante Perez\n"
        "Romilson Lemes Cordeiro\n"
        "Jason Niu\n"
        "Jack Si"
    )

    doc = setup_document(
        "Phase 1 — Project Planning",
        "Hand-Gesture Controlled Fruit-Slicing Game",
        group,
    )

    # ── 1. Scenario Selected ────────────────────────────────────────────
    doc.add_heading("1. Scenario Selected", level=1)
    doc.add_paragraph(
        "This project follows a custom scenario approved by the instructor. "
        "Rather than working with a static medical imaging classification task, "
        "our team has chosen to develop a real-time, interactive fruit-slicing game "
        "that leverages computer vision and hand-gesture recognition as the primary "
        "input method. The project applies core computer vision concepts — including "
        "image processing, object detection, landmark tracking, and real-time "
        "inference — to a practical, engaging application."
    )

    # ── 2. Problem Definition ───────────────────────────────────────────
    doc.add_heading("2. Problem Definition", level=1)
    doc.add_paragraph(
        "Motion-controlled gaming has historically been confined to expensive, "
        "proprietary hardware ecosystems such as the Microsoft Kinect, PlayStation "
        "Move, and modern VR controllers. These systems require specialized sensors "
        "and devices that are not accessible to the average consumer, creating a "
        "barrier to entry for interactive, gesture-based entertainment."
    )
    doc.add_paragraph(
        "At the same time, advances in computer vision — particularly lightweight "
        "hand-tracking frameworks like Google's MediaPipe — have made real-time hand "
        "landmark detection feasible on standard consumer hardware (laptops with "
        "built-in webcams, low-cost USB cameras). However, the majority of "
        "applications leveraging these models remain academic demonstrations and "
        "proof-of-concept prototypes; very few deliver the responsiveness, polish, "
        "and user experience required for an enjoyable, production-quality game."
    )
    doc.add_paragraph(
        "A critical technical challenge is input latency. Gesture recognition "
        "pipelines introduce processing delays between the physical hand movement "
        "and the corresponding on-screen response. In a fast-paced game such as "
        "fruit slicing — where reaction time is the core mechanic — even small "
        "amounts of lag degrade the player experience significantly. The system must "
        "also handle variable lighting conditions, partial occlusions, and different "
        "hand sizes/shapes to maintain robust tracking across diverse environments."
    )
    doc.add_paragraph(
        "This project aims to bridge the gap between accessible computer vision "
        "technology and a polished, real-time gaming experience by building a "
        "hand-gesture controlled fruit-slicing game that runs entirely in a web "
        "browser using a standard webcam — no specialized hardware required."
    )

    # ── 3. Business Goals ───────────────────────────────────────────────
    doc.add_heading("3. Business Goals", level=1)
    doc.add_paragraph(
        "The following five business goals guide the development and success "
        "criteria for this project:"
    )

    goals = [
        (
            "Real-Time Performance: ",
            "Achieve a hand-tracking and game-rendering pipeline that maintains a "
            "minimum of 30 frames per second (FPS) on standard consumer hardware "
            "(e.g., a laptop with an integrated webcam and no dedicated GPU). The "
            "end-to-end latency from hand movement to on-screen blade response must "
            "not exceed 100 milliseconds to ensure a responsive, enjoyable gameplay "
            "experience."
        ),
        (
            "Precise Gesture Detection: ",
            "Implement a slice-detection algorithm that accurately differentiates "
            "between intentional slicing gestures and idle hand movements. The system "
            "must calculate hand velocity, trajectory, and position to determine "
            "whether a fruit has been sliced, achieving a gesture recognition "
            "accuracy of at least 90% during gameplay."
        ),
        (
            "Hardware Accessibility: ",
            "Design the application so it runs entirely in a web browser with no "
            "additional hardware beyond a standard webcam. The game must be playable "
            "on Windows, macOS, and Linux without requiring software installation, "
            "specialized sensors, or GPU-dependent libraries — lowering the barrier "
            "to entry for the widest possible audience."
        ),
        (
            "Intuitive, Gesture-Only User Interface: ",
            "Build the complete user experience — including menu navigation, game "
            "start, pause, and restart — using hand gestures as the sole input "
            "modality. Players should be able to discover and use controls naturally "
            "within the first 30 seconds of interaction, without the need for a "
            "tutorial or on-screen text instructions."
        ),
        (
            "Extensible Computer Vision Framework: ",
            "Architect the gesture-recognition and game-logic modules in a decoupled, "
            "modular design so that additional gestures, game modes, or visual effects "
            "can be integrated in future development phases. This ensures the project "
            "serves as a reusable foundation for broader gesture-based applications "
            "beyond the initial fruit-slicing game."
        ),
    ]

    for i, (title, description) in enumerate(goals, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{i}. ")
        run_num.bold = True
        run_title = p.add_run(title)
        run_title.bold = True
        p.add_run(description)

    # ── 4. Proposed Data Preparation ────────────────────────────────────
    doc.add_heading("4. Proposed Data Preparation", level=1)
    doc.add_paragraph(
        "Unlike traditional image classification projects that rely on static, "
        "labeled datasets, this project processes live video frames in real-time. "
        "However, the project still involves structured image/data organization "
        "for development, testing, and validation purposes."
    )

    doc.add_heading("4.1 Real-Time Data Pipeline", level=2)
    doc.add_paragraph(
        "The primary input is a continuous video stream captured from the user's "
        "webcam. Each frame is processed through Google MediaPipe's Hand Landmark "
        "model, which extracts 21 three-dimensional landmarks per detected hand. "
        "These landmarks are then fed into the game logic to determine hand position, "
        "velocity, and gesture classification (idle, open palm, swiping)."
    )

    doc.add_heading("4.2 Gesture Dataset (for validation and testing)", level=2)
    doc.add_paragraph(
        "To validate gesture recognition accuracy and robustness, the team will "
        "collect and organize a gesture image dataset with the following folder "
        "structure:"
    )

    # Folder structure as a table
    add_formatted_table(
        doc,
        ["Folder", "Label", "Description", "Target Count"],
        [
            ["gesture_data/open_hand/", "open_hand", "Hand open with fingers spread — idle state", "~200 images"],
            ["gesture_data/closed_fist/", "closed_fist", "Hand closed — grabbing or inactive state", "~200 images"],
            ["gesture_data/swipe_left/", "swipe_left", "Hand moving rapidly leftward — slicing gesture", "~200 images"],
            ["gesture_data/swipe_right/", "swipe_right", "Hand moving rapidly rightward — slicing gesture", "~200 images"],
            ["gesture_data/swipe_down/", "swipe_down", "Hand moving rapidly downward — slicing gesture", "~200 images"],
            ["gesture_data/no_hand/", "no_hand", "Background frames with no hand detected", "~200 images"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "This dataset will be captured from the webcam during development sessions "
        "and used to measure classification precision. Each folder name directly "
        "indicates the gesture label for classification."
    )

    doc.add_heading("4.3 Project Folder Structure", level=2)

    folder_structure = (
        "project_root/\n"
        "├── src/                        # Source code\n"
        "│   ├── index.html              # Main game page\n"
        "│   ├── game.js                 # Game logic (fruit spawning, scoring)\n"
        "│   ├── tracking.js             # Hand tracking & gesture detection\n"
        "│   ├── ui.js                   # Menu & HUD rendering\n"
        "│   └── styles.css              # Visual styling\n"
        "├── assets/                     # Visual assets\n"
        "│   ├── fruits/                 # Fruit sprite images\n"
        "│   ├── effects/                # Slicing & particle effects\n"
        "│   └── sounds/                 # Audio effects\n"
        "├── gesture_data/               # Gesture validation dataset\n"
        "│   ├── open_hand/              # ~200 labeled images\n"
        "│   ├── closed_fist/            # ~200 labeled images\n"
        "│   ├── swipe_left/             # ~200 labeled images\n"
        "│   ├── swipe_right/            # ~200 labeled images\n"
        "│   ├── swipe_down/             # ~200 labeled images\n"
        "│   └── no_hand/                # ~200 background images\n"
        "├── docs/                       # Documentation\n"
        "│   └── Phase1_Project_Plan.docx\n"
        "└── README.md"
    )

    p = doc.add_paragraph()
    run = p.add_run(folder_structure)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # ── 5. Team & Responsibilities ──────────────────────────────────────
    doc.add_heading("5. Team Roles and Responsibilities", level=1)

    add_formatted_table(
        doc,
        ["Member", "Role", "Responsibilities"],
        [
            ["Angel Daniel Bustamante Perez", "Lead Developer", "Core game logic integration, performance optimization, tech stack architecture"],
            ["Romilson Lemes Cordeiro", "CV Specialist", "Hand-landmark mapping, gesture recognition pipeline, slicing physics"],
            ["Jason Niu", "CV Specialist", "Gesture classification, real-time data processing, tracking accuracy tuning"],
            ["Jack Si", "UI/UX & Assets", "Visual asset design (fruits, effects), UI layout, project documentation"],
        ],
    )

    # Save
    path = os.path.join(output_dir, "Phase1_Project_Plan_FruitSlicing.docx")
    doc.save(path)
    print(f"[OK] Created: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════
# VERSION B: Brain Tumor Classification (Scenario 1)
# ═══════════════════════════════════════════════════════════════════════════

def create_version_b(output_dir):
    group = (
        "Group 9\n"
        "Angel Daniel Bustamante Perez\n"
        "Romilson Lemes Cordeiro\n"
        "Jason Niu\n"
        "Jack Si"
    )

    doc = setup_document(
        "Phase 1 — Project Planning",
        "Brain Tumor Classification from MRI Scans",
        group,
    )

    # ── 1. Scenario Selected ────────────────────────────────────────────
    doc.add_heading("1. Scenario Selected", level=1)
    doc.add_paragraph(
        "Scenario 1 — Brain Tumor Classification"
    )
    doc.add_paragraph(
        "Our team is part of a healthcare research initiative tasked with assisting "
        "radiologists in diagnosing brain tumours. We have access to a comprehensive "
        "brain MRI dataset containing thousands of images, each labelled as either "
        "\"yes\" (tumour present) or \"no\" (no tumour detected). Our objective is to "
        "build a robust Convolutional Neural Network (CNN) image classifier that can "
        "accurately determine the presence of brain tumours from MRI scans, helping "
        "to speed up diagnosis and improve accuracy for early detection."
    )
    p = doc.add_paragraph("Data Source: ")
    run = p.runs[0]
    run.bold = True
    p.add_run(
        "Brain Tumour Dataset — "
        "https://www.kaggle.com/datasets/ihtishamkhattak99/brain-tumor-dataset-yes-no"
    )

    # ── 2. Problem Definition ───────────────────────────────────────────
    doc.add_heading("2. Problem Definition", level=1)
    doc.add_paragraph(
        "Brain tumours are among the most serious and life-threatening medical "
        "conditions worldwide. According to the World Health Organization (WHO), "
        "central nervous system tumours account for approximately 3% of all cancer "
        "diagnoses, with gliomas being the most common type. Early and accurate "
        "detection is critical — delayed diagnosis often leads to tumour progression, "
        "reduced treatment options, and significantly lower survival rates."
    )
    doc.add_paragraph(
        "The standard diagnostic workflow relies heavily on radiologists manually "
        "reviewing Magnetic Resonance Imaging (MRI) scans. This process presents "
        "several challenges:"
    )

    challenges = [
        (
            "Volume Overload: ",
            "Radiology departments routinely process hundreds of MRI scans per day. "
            "The sheer volume of images can lead to reviewer fatigue, increasing the "
            "risk of missed or delayed diagnoses."
        ),
        (
            "Subjectivity: ",
            "Tumour identification from MRI scans requires expertise and is inherently "
            "subjective. Inter-observer variability between radiologists can lead to "
            "inconsistent diagnoses, particularly for small or early-stage tumours."
        ),
        (
            "Time Pressure: ",
            "In emergency settings, patients may wait hours or days for a specialist "
            "review. An automated screening tool could provide immediate preliminary "
            "results, enabling faster triage and prioritization of critical cases."
        ),
        (
            "Resource Inequality: ",
            "Many healthcare systems — particularly in rural and under-resourced "
            "regions — lack access to specialized neuroradiologists. An AI classifier "
            "could help bridge this gap by providing a reliable second opinion."
        ),
    ]

    for title, desc in challenges:
        add_bullet(doc, desc, bold_prefix=title)

    doc.add_paragraph(
        "This project aims to develop a CNN-based image classifier trained on a "
        "labelled brain MRI dataset that can accurately differentiate between scans "
        "with and without tumours. The classifier is intended to serve as a decision-"
        "support tool — not a replacement for radiologists — providing a fast, "
        "consistent, and objective initial screening to complement expert review."
    )

    # ── 3. Business Goals ───────────────────────────────────────────────
    doc.add_heading("3. Business Goals", level=1)
    doc.add_paragraph(
        "The following five business goals define the success criteria and intended "
        "impact of the brain tumour classification system:"
    )

    goals = [
        (
            "High Diagnostic Accuracy: ",
            "Train a CNN model that achieves a classification accuracy of at least "
            "92% on the held-out test set, with particular emphasis on minimizing "
            "false negatives (missed tumours). The model must demonstrate high "
            "sensitivity (recall ≥ 95%) to ensure that tumour-positive cases are "
            "reliably flagged for expert review."
        ),
        (
            "Reduced Diagnosis Turnaround Time: ",
            "Provide an automated preliminary screening result within seconds of "
            "receiving an MRI scan, compared to the current average turnaround time "
            "of hours or days for manual radiologist review. This enables faster "
            "triage and prioritization of patients who require urgent attention."
        ),
        (
            "Support for Early Detection: ",
            "Enhance the ability to detect brain tumours at earlier stages by "
            "identifying subtle imaging patterns that may be overlooked during "
            "manual review. Early detection directly correlates with improved "
            "treatment outcomes and higher five-year survival rates."
        ),
        (
            "Radiologist Workflow Integration: ",
            "Design the classifier as a complementary decision-support tool that "
            "integrates into existing radiology workflows. The system should present "
            "results in a clear, interpretable format — including confidence scores "
            "and highlighted regions of interest — so radiologists can quickly "
            "validate or override the AI's assessment."
        ),
        (
            "Scalability and Accessibility: ",
            "Build the model using open-source frameworks (TensorFlow/Keras or "
            "PyTorch) so that it can be deployed at low cost in diverse healthcare "
            "settings, including resource-constrained environments. The goal is to "
            "make high-quality AI-assisted screening accessible beyond major urban "
            "medical centres."
        ),
    ]

    for i, (title, description) in enumerate(goals, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{i}. ")
        run_num.bold = True
        run_title = p.add_run(title)
        run_title.bold = True
        p.add_run(description)

    # ── 4. Proposed Data Preparation ────────────────────────────────────
    doc.add_heading("4. Proposed Data Preparation", level=1)

    doc.add_heading("4.1 Dataset Overview", level=2)
    doc.add_paragraph(
        "The project uses the Brain Tumor Dataset available on Kaggle, created by "
        "Ihtisham Khattak. The dataset contains MRI scan images organized into two "
        "classes for binary classification."
    )

    add_formatted_table(
        doc,
        ["Property", "Details"],
        [
            ["Source", "Kaggle — Brain Tumor Dataset (Yes/No)"],
            ["Total Images", "~3,000 MRI scan images"],
            ["Classes", "2 (Yes — tumour present, No — no tumour)"],
            ["Image Format", "JPEG (.jpg)"],
            ["Image Type", "Grayscale brain MRI cross-sections"],
            ["Labels", "Supervised — folder names serve as class labels"],
        ],
    )

    doc.add_heading("4.2 Folder Structure", level=2)
    doc.add_paragraph(
        "The images are organized into clearly labelled folders that directly "
        "indicate the classification labels. The proposed folder structure for "
        "the project is as follows:"
    )

    folder_structure = (
        "brain_tumor_dataset/\n"
        "├── yes/                    # MRI images WITH brain tumour\n"
        "│   ├── Y1.jpg\n"
        "│   ├── Y2.jpg\n"
        "│   ├── Y3.jpg\n"
        "│   └── ... (~1,500 images)\n"
        "├── no/                     # MRI images WITHOUT brain tumour\n"
        "│   ├── 1 no.jpg\n"
        "│   ├── 2 no.jpg\n"
        "│   ├── 3 no.jpg\n"
        "│   └── ... (~1,500 images)\n"
        "└── README.md               # Dataset documentation"
    )

    p = doc.add_paragraph()
    run = p.add_run(folder_structure)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_paragraph(
        "The folder names \"yes\" and \"no\" serve as the ground-truth labels "
        "for supervised training. This two-folder structure is the standard "
        "convention for binary image classification in frameworks like "
        "TensorFlow/Keras (via ImageDataGenerator or tf.data) and PyTorch "
        "(via ImageFolder)."
    )

    doc.add_heading("4.3 Data Preprocessing Steps", level=2)
    doc.add_paragraph(
        "The following preprocessing steps will be applied to prepare the "
        "images for model training:"
    )

    steps = [
        (
            "Resizing: ",
            "All images will be resized to a uniform resolution of 224×224 pixels "
            "to match the input requirements of standard CNN architectures (e.g., "
            "VGG16, ResNet50)."
        ),
        (
            "Normalization: ",
            "Pixel values will be scaled to the range [0, 1] by dividing by 255, "
            "ensuring consistent input magnitudes across all images."
        ),
        (
            "Train/Validation/Test Split: ",
            "The dataset will be split into 70% training, 15% validation, and 15% "
            "test sets using stratified sampling to maintain class balance."
        ),
        (
            "Data Augmentation: ",
            "To increase robustness and reduce overfitting, augmentation techniques "
            "will be applied to the training set, including random horizontal flips, "
            "rotations (±15°), zoom (±10%), and brightness adjustments."
        ),
    ]

    for title, desc in steps:
        add_bullet(doc, desc, bold_prefix=title)

    doc.add_heading("4.4 Complete Project Folder Structure", level=2)

    proj_structure = (
        "project_root/\n"
        "├── brain_tumor_dataset/          # Raw image dataset\n"
        "│   ├── yes/                      # Tumour-positive MRI images\n"
        "│   └── no/                       # Tumour-negative MRI images\n"
        "├── notebooks/                    # Jupyter notebooks\n"
        "│   ├── 01_EDA.ipynb              # Exploratory data analysis\n"
        "│   ├── 02_Preprocessing.ipynb    # Data preparation pipeline\n"
        "│   └── 03_Model_Training.ipynb   # Model training & evaluation\n"
        "├── src/                          # Source code modules\n"
        "│   ├── data_loader.py            # Dataset loading utilities\n"
        "│   ├── model.py                  # CNN architecture definition\n"
        "│   ├── train.py                  # Training loop\n"
        "│   └── evaluate.py              # Evaluation & metrics\n"
        "├── models/                       # Saved model weights\n"
        "├── docs/                         # Documentation\n"
        "│   └── Phase1_Project_Plan.docx\n"
        "├── requirements.txt              # Python dependencies\n"
        "└── README.md"
    )

    p = doc.add_paragraph()
    run = p.add_run(proj_structure)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # ── 5. Team & Responsibilities ──────────────────────────────────────
    doc.add_heading("5. Team Roles and Responsibilities", level=1)

    add_formatted_table(
        doc,
        ["Member", "Role", "Responsibilities"],
        [
            ["Angel Daniel Bustamante Perez", "Lead Developer", "CNN architecture design, model training pipeline, hyperparameter tuning"],
            ["Romilson Lemes Cordeiro", "Data Engineer", "Dataset acquisition, preprocessing pipeline, data augmentation, train/val/test splitting"],
            ["Jason Niu", "ML Engineer", "Model evaluation, metrics analysis (accuracy, recall, F1), confusion matrix visualization"],
            ["Jack Si", "Project Manager", "Documentation, project planning, results presentation, Phase 1 & 2 report preparation"],
        ],
    )

    # Save
    path = os.path.join(output_dir, "Phase1_Project_Plan_BrainTumor.docx")
    doc.save(path)
    print(f"[OK] Created: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════
# VERSION C: Lung Cancer Classification (Scenario 2)
# ═══════════════════════════════════════════════════════════════════════════

def create_version_c(output_dir):
    group = (
        "Group 9\n"
        "Angel Daniel Bustamante Perez\n"
        "Romilson Lemes Cordeiro\n"
        "Jason Niu\n"
        "Jack Si"
    )

    doc = setup_document(
        "Phase 1 \u2014 Project Planning",
        "Lung Cancer Classification from CT Scans",
        group,
    )

    # ── 1. Scenario Selected ────────────────────────────────────────────
    doc.add_heading("1. Scenario Selected", level=1)
    doc.add_paragraph(
        "Scenario 2 \u2014 Lung Cancer Classification"
    )
    doc.add_paragraph(
        "As researchers affiliated with the Iraq-Oncology Teaching Hospital / "
        "National Center for Cancer Diseases (IQ-OTH/NCCD), our team has been "
        "tasked with developing an image classifier that can accurately "
        "differentiate between normal lung tissue, benign lung tumours, and "
        "malignant lung tumours using CT scan images. The dataset, collected in "
        "the fall of 2019, includes 1,190 CT scan images from 110 cases, "
        "carefully annotated by expert oncologists and radiologists."
    )
    p = doc.add_paragraph("Data Source: ")
    run = p.runs[0]
    run.bold = True
    p.add_run(
        "IQ-OTH/NCCD Lung Cancer Dataset \u2014 "
        "https://www.kaggle.com/datasets/adityamahimkar/iqothnccd-lung-cancer-dataset"
    )

    # ── 2. Problem Definition ───────────────────────────────────────────
    doc.add_heading("2. Problem Definition", level=1)
    doc.add_paragraph(
        "Lung cancer is the leading cause of cancer-related deaths worldwide, "
        "accounting for approximately 1.8 million fatalities annually according "
        "to the World Health Organization (WHO). The five-year survival rate for "
        "lung cancer is heavily dependent on the stage at which it is diagnosed: "
        "patients identified at an early, localized stage have a survival rate "
        "exceeding 60%, while late-stage diagnoses drop this figure to below 10%. "
        "This stark contrast underscores the critical importance of early and "
        "accurate detection."
    )
    doc.add_paragraph(
        "Computed Tomography (CT) scans are the primary imaging modality used "
        "for lung cancer screening and diagnosis. Radiologists must carefully "
        "analyse cross-sectional images of the lungs to identify abnormalities "
        "and classify them as normal tissue, benign growths, or malignant tumours. "
        "This process presents several significant challenges:"
    )

    challenges = [
        (
            "Diagnostic Complexity: ",
            "Unlike binary classification tasks, lung cancer diagnosis requires "
            "distinguishing between three classes \u2014 normal tissue, benign tumours, "
            "and malignant tumours. Benign and malignant growths can appear visually "
            "similar on CT scans, making differentiation highly challenging even for "
            "experienced radiologists."
        ),
        (
            "Inter-Observer Variability: ",
            "Studies have shown that radiologist agreement on lung nodule "
            "classification can vary significantly, particularly for small or "
            "ambiguous lesions. This subjectivity means that the same scan may "
            "receive different interpretations depending on the reviewing physician, "
            "leading to inconsistent patient outcomes."
        ),
        (
            "Screening Volume: ",
            "With increased adoption of low-dose CT screening programmes for "
            "high-risk populations (e.g., long-term smokers over age 50), "
            "radiology departments face a growing volume of scans that must be "
            "reviewed. This workload increases the risk of fatigue-related errors "
            "and delays in reporting."
        ),
        (
            "Resource Constraints: ",
            "Many healthcare facilities \u2014 particularly in developing nations and "
            "rural areas \u2014 lack access to specialist thoracic radiologists. An "
            "AI-assisted classification tool could provide a reliable preliminary "
            "assessment, enabling faster triage and reducing dependency on scarce "
            "specialist resources."
        ),
        (
            "False Positive Burden: ",
            "Lung cancer screening programmes are known to produce a high rate of "
            "false positives, leading to unnecessary invasive procedures such as "
            "biopsies. A more accurate classifier could help reduce false positive "
            "rates, decreasing patient anxiety and healthcare costs associated with "
            "follow-up investigations."
        ),
    ]

    for title, desc in challenges:
        add_bullet(doc, desc, bold_prefix=title)

    doc.add_paragraph(
        "This project aims to develop a deep learning image classifier \u2014 based "
        "on Convolutional Neural Networks (CNNs) \u2014 trained on the IQ-OTH/NCCD "
        "lung cancer CT dataset. The model will perform three-class classification "
        "(normal, benign, malignant) to assist radiologists by providing a fast, "
        "objective, and consistent preliminary assessment that complements their "
        "expert judgment."
    )

    # ── 3. Business Goals ───────────────────────────────────────────────
    doc.add_heading("3. Business Goals", level=1)
    doc.add_paragraph(
        "The following five business goals define the success criteria and "
        "intended impact of the lung cancer classification system:"
    )

    goals = [
        (
            "High Multi-Class Accuracy: ",
            "Train a CNN model that achieves an overall classification accuracy "
            "of at least 90% across all three classes (normal, benign, malignant) "
            "on the held-out test set. The model must demonstrate particularly high "
            "sensitivity (recall >= 93%) for the malignant class to minimize the risk "
            "of missed cancer diagnoses, which have the most severe clinical "
            "consequences."
        ),
        (
            "Reduction in False Positive Rates: ",
            "Reduce the false positive rate for benign and malignant classifications "
            "by providing a more objective and consistent assessment compared to "
            "manual review alone. The target is to achieve a specificity of at least "
            "88% for the normal class, reducing unnecessary follow-up procedures "
            "and the associated patient burden."
        ),
        (
            "Accelerated Screening Workflow: ",
            "Enable radiologists to process CT scan reviews more efficiently by "
            "providing an automated preliminary classification within seconds. "
            "The system should reduce the average time-to-initial-assessment from "
            "hours (in high-volume departments) to under 10 seconds per scan, "
            "allowing specialists to focus their attention on flagged abnormal cases."
        ),
        (
            "Clinical Decision Support with Interpretability: ",
            "Present classification results alongside confidence scores and, where "
            "feasible, visual heatmaps (e.g., using Grad-CAM) highlighting the "
            "regions of the CT scan that most influenced the model's decision. This "
            "interpretability is essential for building clinician trust and ensuring "
            "the tool integrates effectively into diagnostic workflows rather than "
            "operating as a black box."
        ),
        (
            "Scalable and Cost-Effective Deployment: ",
            "Build the classification model using open-source deep learning "
            "frameworks (TensorFlow/Keras or PyTorch) with the capability to run "
            "inference on standard hospital computing infrastructure (CPU-based "
            "servers or modest GPU hardware). The solution must be deployable "
            "without requiring expensive, specialized AI hardware, making it "
            "accessible to healthcare facilities in both developed and developing "
            "regions."
        ),
    ]

    for i, (title, description) in enumerate(goals, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{i}. ")
        run_num.bold = True
        run_title = p.add_run(title)
        run_title.bold = True
        p.add_run(description)

    # ── 4. Proposed Data Preparation ────────────────────────────────────
    doc.add_heading("4. Proposed Data Preparation", level=1)

    doc.add_heading("4.1 Dataset Overview", level=2)
    doc.add_paragraph(
        "The project uses the IQ-OTH/NCCD Lung Cancer Dataset available on "
        "Kaggle. The dataset was collected at the Iraq-Oncology Teaching Hospital "
        "and the National Center for Cancer Diseases in the fall of 2019. All "
        "images were annotated by expert oncologists and radiologists."
    )

    add_formatted_table(
        doc,
        ["Property", "Details"],
        [
            ["Source", "Kaggle \u2014 IQ-OTH/NCCD Lung Cancer Dataset"],
            ["Total Images", "1,190 CT scan images"],
            ["Cases", "110 patient cases"],
            ["Classes", "3 (Normal, Benign, Malignant)"],
            ["Image Format", "JPEG (.jpg) / PNG (.png)"],
            ["Image Type", "Grayscale or RGB chest CT cross-sections"],
            ["Annotation", "Supervised \u2014 expert oncologist and radiologist labels"],
            ["Collection Period", "Fall 2019"],
        ],
    )

    doc.add_heading("4.2 Folder Structure", level=2)
    doc.add_paragraph(
        "The images are organized into three clearly labelled folders that "
        "directly correspond to the classification labels. The proposed folder "
        "structure for the image dataset is as follows:"
    )

    folder_structure = (
        "lung_cancer_dataset/\n"
        "\u251c\u2500\u2500 normal/                   # CT images of normal lung tissue\n"
        "\u2502   \u251c\u2500\u2500 normal_001.jpg\n"
        "\u2502   \u251c\u2500\u2500 normal_002.jpg\n"
        "\u2502   \u2514\u2500\u2500 ... (~400 images)\n"
        "\u251c\u2500\u2500 benign/                   # CT images of benign lung tumours\n"
        "\u2502   \u251c\u2500\u2500 benign_001.jpg\n"
        "\u2502   \u251c\u2500\u2500 benign_002.jpg\n"
        "\u2502   \u2514\u2500\u2500 ... (~400 images)\n"
        "\u251c\u2500\u2500 malignant/                # CT images of malignant lung tumours\n"
        "\u2502   \u251c\u2500\u2500 malignant_001.jpg\n"
        "\u2502   \u251c\u2500\u2500 malignant_002.jpg\n"
        "\u2502   \u2514\u2500\u2500 ... (~390 images)\n"
        "\u2514\u2500\u2500 README.md                 # Dataset documentation"
    )

    p = doc.add_paragraph()
    run = p.add_run(folder_structure)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_paragraph(
        "The folder names \"normal\", \"benign\", and \"malignant\" serve as the "
        "ground-truth labels for supervised training. This multi-folder structure "
        "is the standard convention for multi-class image classification in "
        "frameworks like TensorFlow/Keras (via image_dataset_from_directory) and "
        "PyTorch (via ImageFolder), where each subfolder automatically maps to a "
        "distinct class label."
    )

    doc.add_heading("4.3 Data Preprocessing Steps", level=2)
    doc.add_paragraph(
        "The following preprocessing steps will be applied to prepare the "
        "CT scan images for model training:"
    )

    steps = [
        (
            "Resizing: ",
            "All images will be resized to a uniform resolution of 224x224 pixels "
            "to match the input requirements of standard CNN architectures such as "
            "VGG16, ResNet50, and EfficientNet."
        ),
        (
            "Normalization: ",
            "Pixel values will be scaled to the range [0, 1] by dividing by 255. "
            "For transfer learning with pre-trained ImageNet models, channel-wise "
            "mean subtraction will also be applied."
        ),
        (
            "Train/Validation/Test Split: ",
            "The dataset will be split into 70% training, 15% validation, and 15% "
            "test sets using stratified sampling to ensure that all three classes "
            "(normal, benign, malignant) are proportionally represented in each "
            "subset."
        ),
        (
            "Data Augmentation: ",
            "To address the relatively small dataset size (1,190 images) and reduce "
            "overfitting, augmentation techniques will be applied to the training "
            "set. These include random horizontal/vertical flips, rotations "
            "(ranging from -20 degrees to +20 degrees), zoom (ranging from -15% to +15%), brightness and "
            "contrast adjustments, and elastic deformations."
        ),
        (
            "Class Balancing: ",
            "If class distributions are imbalanced, techniques such as oversampling "
            "the minority class, undersampling the majority class, or applying "
            "class-weight adjustments during training will be used to prevent the "
            "model from being biased toward the most frequent class."
        ),
    ]

    for title, desc in steps:
        add_bullet(doc, desc, bold_prefix=title)

    doc.add_heading("4.4 Complete Project Folder Structure", level=2)

    proj_structure = (
        "project_root/\n"
        "\u251c\u2500\u2500 lung_cancer_dataset/            # Raw image dataset\n"
        "\u2502   \u251c\u2500\u2500 normal/                     # Normal lung tissue CT images\n"
        "\u2502   \u251c\u2500\u2500 benign/                     # Benign lung tumour CT images\n"
        "\u2502   \u2514\u2500\u2500 malignant/                  # Malignant lung tumour CT images\n"
        "\u251c\u2500\u2500 notebooks/                      # Jupyter notebooks\n"
        "\u2502   \u251c\u2500\u2500 01_EDA.ipynb                # Exploratory data analysis\n"
        "\u2502   \u251c\u2500\u2500 02_Preprocessing.ipynb      # Data preparation pipeline\n"
        "\u2502   \u2514\u2500\u2500 03_Model_Training.ipynb     # Model training & evaluation\n"
        "\u251c\u2500\u2500 src/                            # Source code modules\n"
        "\u2502   \u251c\u2500\u2500 data_loader.py              # Dataset loading utilities\n"
        "\u2502   \u251c\u2500\u2500 model.py                    # CNN architecture definition\n"
        "\u2502   \u251c\u2500\u2500 train.py                    # Training loop\n"
        "\u2502   \u2514\u2500\u2500 evaluate.py                 # Evaluation & metrics\n"
        "\u251c\u2500\u2500 models/                         # Saved model weights\n"
        "\u251c\u2500\u2500 docs/                           # Documentation\n"
        "\u2502   \u2514\u2500\u2500 Phase1_Project_Plan.docx\n"
        "\u251c\u2500\u2500 requirements.txt                # Python dependencies\n"
        "\u2514\u2500\u2500 README.md"
    )

    p = doc.add_paragraph()
    run = p.add_run(proj_structure)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # ── 5. Team & Responsibilities ──────────────────────────────────────
    doc.add_heading("5. Team Roles and Responsibilities", level=1)

    add_formatted_table(
        doc,
        ["Member", "Role", "Responsibilities"],
        [
            ["Angel Daniel Bustamante Perez", "Lead Developer",
             "CNN architecture design (multi-class), model training pipeline, "
             "hyperparameter tuning, transfer learning implementation"],
            ["Romilson Lemes Cordeiro", "Data Engineer",
             "Dataset acquisition from Kaggle, preprocessing pipeline, data "
             "augmentation, train/val/test splitting, class balancing"],
            ["Jason Niu", "ML Engineer",
             "Model evaluation, multi-class metrics analysis (accuracy, per-class "
             "recall/precision, F1, confusion matrix), Grad-CAM visualization"],
            ["Jack Si", "Project Manager",
             "Documentation, project planning, results presentation, Phase 1 & 2 "
             "report preparation, team coordination"],
        ],
    )

    # Save
    path = os.path.join(output_dir, "Phase1_Project_Plan_LungCancer.docx")
    doc.save(path)
    print(f"[OK] Created: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    output_dir = os.path.dirname(os.path.abspath(__file__))
    print("Generating Phase 1 documents...\n")

    path_a = create_version_a(output_dir)
    path_b = create_version_b(output_dir)
    path_c = create_version_c(output_dir)

    print(f"\n[OK] All three documents generated in:\n   {output_dir}")
